from flask import Flask, render_template, request, send_file
from docx import Document
import io
import os
import spire.doc as spd
import fitz  # PyMuPDF
app = Flask(__name__)

@app.route("/")
def index():
    return render_template('index.html')

@app.route('/generate_quote', methods=['POST'])
def generate_quote():
    

    def remove_evaluation_warning(pdf_path, output_path):
        pdf_document = fitz.open(pdf_path)
    
    # Iterate through the pages
        for page_num in range(len(pdf_document)):
            page = pdf_document.load_page(page_num)
            text_instances = page.search_for("Evaluation Warning: The document was created with Spire.Doc for Python.")

            # Remove the specific text occurrences
            for inst in text_instances:
                page.add_redact_annot(inst)  # Mark the text for redaction
            
            # Apply the redactions (removes the marked text)
            page.apply_redactions()

        # Save the result to a new PDF
        pdf_document.save(output_path)
        pdf_document.close()

        



    name = request.form['name']
    address1 = request.form['address1']
    address2 = request.form['address2']
    source = request.form['source']
    destination = request.form['destination']
    amount = request.form['amount']
    date = request.form['date']
    notes = request.form['notes']
    phone = request.form['phone']
    email = request.form['email']
    template = request.form['template']

    # Combine address lines
    address = f"{address1}\n{address2}"

    # Choose the template based on the selected option
    if template == 'AUM':
        template_file = 'QuotationTemp_AUM.docx'
    else:
        template_file = 'QuotationTemp_International.docx'

    # Load the docx template
    doc = Document(template_file)

    # Replace placeholders with actual values
    for paragraph in doc.paragraphs:
        if '${name}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${name}', name)
        if '${address}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${address}', address)
        if '${source}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${source}', source)
        if '${destination}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${destination}', destination)
        if '${amount}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${amount}', amount)
        if '${date}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${date}', date)
        if '${notes}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${notes}', notes)
        if '${phone}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${phone}', phone)
        if '${email}' in paragraph.text:
            paragraph.text = paragraph.text.replace('${email}', email)

    # Save the modified document to a temporary file
    temp_docx = 'temp_quote.docx'
    doc.save(temp_docx)

    # Convert DOCX to PDF using Spire.Doc
    pdf_file = 'quote.pdf'
    document = spd.Document()
    document.LoadFromFile(temp_docx)
    document.SaveToFile(pdf_file, spd.FileFormat.PDF)
    document.Close()
    # Example usage
    input_pdf = "quote.pdf"
    output_pdf = "quote_no_warning.pdf"
    remove_evaluation_warning(input_pdf, output_pdf)

    # Read the PDF file into a BytesIO object
    buffer = io.BytesIO()
    with open(output_pdf, 'rb') as f:
        buffer.write(f.read())
    buffer.seek(0)

    # Clean up temporary files
    os.remove(temp_docx)
    os.remove(pdf_file)
    os.remove(output_pdf)

    return send_file(buffer, as_attachment=True, download_name='quote.pdf', mimetype='application/pdf')

if __name__ == "__main__":
    app.run(debug=True)
